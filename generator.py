import itertools
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageDraw, ImageFont
import sys

file = docx.Document() # Создание word-файла
style = file.styles["Normal"] # Текстовый стиль word-файла
style.font.name = "Liberation Serif" # Используемый шрифт
style.font.size = Pt(14) # Размер шрифта
Serial_Number = 1 # Номер соответствия

# Создание рисунка соответствия
def make_draw(Q, X, Y):
    try:
        image = Image.open("unfill.jpg")
    except:
        print("Unable to load image")
        sys.exit(1)

    idraw = ImageDraw.Draw(image)
    font = ImageFont.truetype("arial.ttf", size=12)
    shift = 0
    arr_X = dict()
    arr_Y = dict()
    for i in range(len(X)):
        idraw.ellipse((125, 30 + shift, 145, 50 + shift), fill="black")
        arr_X[X[i]] = ((135, 37 + shift))
        idraw.text((105, 25 + shift), text=f"{X[i]}", font=font)
        shift += 30

    shift = 0
    for i in range(len(Y)):
        idraw.ellipse((260, 30 + shift, 280, 50 + shift), fill="black")
        arr_Y[Y[i]] = ((270, 37 + shift))
        idraw.text((285, 25 + shift), text=f"{Y[i]}", font=font)
        shift += 30

    for i in range(len(Q)):
        idraw.line((arr_X[Q[i][0]], arr_Y[Q[i][1]]), fill="black", width=3)
    image.save('image.png')
def bold(text):
    q = file.add_paragraph()
    p = q.add_run(text)
    p.bold = True

    return q

def head(text):
    p = file.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(16)

def head_with_bold(text):
    p = file.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(18)
    r.bold = True
def handling(Q: list, X, Y) -> None:
    global Serial_Number
    Q = set(Q)
    G = dict()
    anti_G = dict()
    D = set()
    Im = set()

    for el in Q:
        # Fill G - X set
        if el[0] in G:
            G[el[0]].add(el[1])
        else:
            G[el[0]] = {el[1]}

        # Fill anti_G - Y set
        if el[1] in anti_G:
            anti_G[el[1]].add(el[0])
        else:
            anti_G[el[1]] = {el[0]}

        D.add(el[0])
        Im.add(el[1])

    # Создание изображения
    make_draw(list(Q), list(X), list(Y))

    head_with_bold(f"Соответствие Q({Serial_Number})")
    if Q != set():
        file.add_paragraph(f"1) Q{Serial_Number} = {Q}")
    else:
        file.add_paragraph(f"1) Q{Serial_Number} =" + " {}")
    file.add_picture("image.png")
    p = file.add_paragraph("\n2)")
    p = p.add_run("Образы:")
    p.bold = True
    keys = list(G.keys())
    for i in range(len(G)):
        file.add_paragraph(f"G({keys[i]}) = {G[keys[i]]}")
    bold("Прообразы:")
    keys = list(anti_G.keys())
    for i in range(len(anti_G)):
        file.add_paragraph(f"G^-1({keys[i]}) = {anti_G[keys[i]]}")

    # Область определений
    p = bold("Область определений: ")
    if D != set():
        p.add_run(f"D(Q{Serial_Number}) = {D}")
    else:
        p.add_run(f"D(Q{Serial_Number}) =" + " {}")

    # Область значений
    p = bold(f"Область значений: ")
    if Im != set():
        p.add_run(f"Im(Q{Serial_Number}) = {Im}")
    else:
        p.add_run(f"Im(Q{Serial_Number}) =" + " {}")

    file.add_paragraph().add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
    p = file.add_paragraph("\n3)")
    p = p.add_run("Анализ с обоснованием свойств:")
    p.bold = True
    head("\n1.Соответствие:")

    flag_allgod = 0
    # Всюду определенное
    if D == set(X):
        file.add_paragraph(f"1.1) Всюду определенное, т.к область определения D(Q{Serial_Number}) равна X")
        flag_allgod = 1
    else:
        file.add_paragraph(f"1.1) Не всюду определенная, т.к. область определения D(Q{Serial_Number}) не равна X")

    # Сюръективное
    if Im == set(Y):
        file.add_paragraph(f"1.2) Сюръективное, т.к. область значений Im(Q{Serial_Number}) равна Y")
    else:
        file.add_paragraph(f"1.2) Несюръективное, т.к. область значений Im(Q{Serial_Number}) не равна Y")

    flag_func = 0
    flag_ineqt = 0
    # Функциональное
    if all(len(G[x]) <= 1 for x in G):
        file.add_paragraph("1.3) Функциональное, т.к. образ любого элемента множества X содержит не более одного элемента из Y")
        flag_func = 1
    else:
        for x in G:
            if len(G[x]) > 1:
                file.add_paragraph(f"1.3) Нефункцинальное, т.к. образ элемента \"{x}\",\nравный G({x}) = {G[x]} содержит больше одного элемента")
                break

    # Инъективное
    if all(len(anti_G[x]) <= 1 for x in anti_G):
        file.add_paragraph("1.4) Инъективное, т.к. прообраз любого элемента Y содержит не более одного элемента из X")
        flag_ineqt = 1
    else:
        for x in anti_G:
            if len(anti_G[x]) > 1:
                file.add_paragraph(f"1.4) Неинъективное, т.к. прообораз элемента \"{x}\",\nравный Im({x}) = {anti_G[x]} содержит больше одного элемента")
                break

    # Взаимнооднозначное
    if flag_ineqt and flag_func:
        file.add_paragraph("1.5) Взаимнооднозначное, т.к. функционально и инъективно")
    elif flag_ineqt == 1 and flag_func == 0:
        file.add_paragraph("1.5) Невзаимнооднозначное, т.к. нефункционально")
    elif flag_ineqt == 0 and flag_func == 1:
        file.add_paragraph("1.5) Невзаимнооднозначное, т.к. неинъективно")
    else:
        file.add_paragraph("1.5) Невзаимнооднозначное, т.к. нефункционально и неинъективно")

    # Отображение
    head("\n2.Отображение:")

    # Является ли?
    flag_image = 0
    if D == set(X):
        file.add_paragraph(f"2.1) Является отображением, т.к. область определения D(Q{Serial_Number}) равна X")
        flag_image = 1
    else:
        file.add_paragraph(f"2.1) Не является отображением, т.к. область определения D(Q{Serial_Number}) не равна X")

    flag_image_ineqt = 0
    flag_image_sureqt = 0
    if flag_image:
        # Инъективное
        if all(len(anti_G[x]) <= 1 for x in anti_G):
            file.add_paragraph("2.2) Инъективное, т.к. все элементы из Y участвуют в парах не более одного раза")
            flag_image_ineqt = 1
        else:
            for x in anti_G:
                if len(anti_G[x]) > 1:
                    file.add_paragraph(f"2.2) Неинъективное, т.к. элемент \"{x}\" из Y участвует в паре {len(anti_G[x])} раза")
                    break

        # Сюръективное
        if Im == set(Y):
            file.add_paragraph("2.3) Сюръективное, т.к. каждому элементу из Y соответствует прообраз из X")
            flag_image_sureqt = 1
        else:
            for x in Y:
                if not(x in Im):
                    file.add_paragraph(f"2.3) Несюръективное, т.к. элементу \"{x}\" из Y не соответствует элемент из X")

        # Биективное
        if flag_image_sureqt and flag_image_ineqt:
            file.add_paragraph("2.4) Биективное, т.к. инъективное и сюръективное")
        elif flag_image_sureqt == 1 and flag_image_ineqt == 0:
            file.add_paragraph("2.4) Небиективное, т.к. неинъективное")
        elif flag_image_sureqt == 0 and flag_image_ineqt == 1:
            file.add_paragraph("2.4) Небиективное, т.к. несюръективное")
        else:
            file.add_paragraph("2.4) Небиективное, т.к. неинъективное и несюръективное")

    # Функция
    head("\n3.Функция:")

    #Является ли?
    flag_iffunc = 0
    if flag_func and flag_allgod:
        file.add_paragraph("3.1) Является функцией, т.к. функционально и всюду определенно")
        flag_iffunc = 1
    elif flag_func == 1 and flag_allgod == 0:
        file.add_paragraph("3.1) Не является функцией, т.к. не всюду определенно")
    elif flag_func == 0 and flag_allgod == 1:
        file.add_paragraph("3.1) Не является функцией, т.к. нефункционально")
    else:
        file.add_paragraph("3.1) Не является функцией, т.к. нефункционально и не всюду определенно")

    if flag_iffunc:
        # Инъективная
        flag_func_ineqt = 0
        if len(X) <= len(Y) and flag_ineqt:
            file.add_paragraph("3.2) Инъективная, т.к. |X| <= |Y| и любой элемент из X имеет не более одного прообраза")
            flag_func_ineqt = 1
        elif not(len(X) <= len(Y)) and flag_ineqt == 1:
            file.add_paragraph("3.2) Неинъективная, т.к. |X| > |Y|")
        elif len(X) <= len(Y) and flag_ineqt == 0:
            for x in anti_G:
                if len(anti_G[x]) > 1:
                    file.add_paragraph(f"3.2) Неинъективная, т.к. пробораз элемента \"{x}\",\nравный Im({x}) = {anti_G[x]} содержит больше одного элемента")
                    break
        else:
            for x in anti_G:
                if len(anti_G[x]) > 1:
                    file.add_paragraph(f"3.2) Неинъективная, т.к. пробораз элемента \"{x}\",\nравный Im({x}) = {anti_G[x]} содержит больше одного элемента, а также |X| > |Y|")
                    break

        # Сюръективная
        flag_func_sureqt = 0
        if len(X) >= len(Y) and flag_image_sureqt:
            file.add_paragraph("3.3) Сюръективная, т.к. |X| >= |Y| и всем элементам из Y соотвествует элемент из X")
            flag_func_sureqt = 1
        elif not(len(X) >= len(Y)) and flag_image_sureqt == 1:
            file.add_paragraph("3.3) Несюръективная, т.к. |X| < |Y|")
        elif len(X) >= len(Y) and flag_image_sureqt == 0:
            for x in Y:
                if not(x in Im):
                    file.add_paragraph(f"3.3) Несюръективное, т.к. элементу \"{x}\" из Y не соответствует элемент из X")
        else:
            for x in Y:
                if not(x in Im):
                    file.add_paragraph(f"3.3) Несюръективное, т.к. элементу \"{x}\" из Y не соответствует элемент из X,\nа также |X| < |Y|")

        # Биективная
        if len(X) == len(Y) and flag_func_sureqt and flag_func_ineqt:
            file.add_paragraph("3.4) Биективная, т.к. и сюръективная и инъективная и |X| = |Y|")
        elif not(len(X) == len(Y)) and flag_func_sureqt == 1 and flag_func_ineqt == 1:
            file.add_paragraph("3.4) Небиективная, т.к. |X| не равен |Y|")
        elif len(X) == len(Y) and flag_func_sureqt == 0 and flag_func_ineqt == 1:
            file.add_paragraph("3.4) Небиективная, т.к. несюръективная")
        elif len(X) == len(Y) and flag_func_sureqt == 1 and flag_func_ineqt == 0:
            file.add_paragraph("3.4) Небиективная, т.к. неинъективная")
        elif not(len(X) == len(Y)) and flag_func_sureqt == 0 and flag_func_ineqt == 1:
            file.add_paragraph("3.4) Небиективная, т.к. |X| не равен |Y| и несюръективная")
        elif not(len(X) == len(Y)) and flag_func_sureqt == 1 and flag_func_ineqt == 0:
            file.add_paragraph("3.4) Небиективная, т.к. |X| не равен |Y| и неинъективная")
        else:
            file.add_paragraph("3.4) Небиективная, т.к. |X| не равен |Y| и несюръективная и неинъективная")
    file.add_paragraph().add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
    Serial_Number += 1

def generator(X: list, Y: list) -> None:

    power_X = len(X)
    power_Y = len(Y)
    pairs_contain = list(itertools.product(X, Y))
    combinations = []

    # 2 * 2
    if power_X * power_Y == 4:
        for i in range(5):
            combinations += list(itertools.combinations(pairs_contain, i))

        for comb in combinations:
            handling(comb, X, Y)
    # 2 * 3
    else:
        for i in range(7):
            combinations += list(itertools.combinations(pairs_contain, i))

        comb_x64 = []
        flag_1 = 1
        flag_2 = 4
        flag_3 = 4
        flag_4 = 3
        flag_5 = 3
        flag_6 = 1
        for comb in combinations:
            if len(comb) == 1 and flag_1:
                comb_x64.append(comb)
                flag_1 -= 1

            elif len(comb) == 2 and flag_2:
                comb_x64.append(comb)
                flag_2 -= 1

            elif len(comb) == 3 and flag_3:
                comb_x64.append(comb)
                flag_3 -= 1

            elif len(comb) == 4 and flag_4:
                comb_x64.append(comb)
                flag_4 -= 1

            elif len(comb) == 5 and flag_5:
                comb_x64.append(comb)
                flag_5 -= 1

            elif len(comb) == 6 and flag_6:
                comb_x64.append(comb)
                flag_6 -= 1

        for comb in comb_x64:
            handling(comb, X, Y)


if __name__ == "__main__":

    print("Введите множество X = ", end="")
    X = list(map(int, input().split()))

    # The correctness of the values X
    if len(X) != 2 and len(X) != 3:
        print("Неправильное значение X")
        exit()

    print("Введите множество Y = ", end="")
    Y = list(map(int, input().split()))

    # The correctness of the values Y
    if len(Y) != 2 and len(Y) != 3:
        print("Неправильное значение Y")
        exit()

    head_with_bold("Задание")
    bold(f"X = {X}")
    bold(f"Y = {Y}\n")

    # Запуск ракеты
    generator(X, Y)
    file.save("generator.docx")
    print("Файл был успешно сгенерирован!")